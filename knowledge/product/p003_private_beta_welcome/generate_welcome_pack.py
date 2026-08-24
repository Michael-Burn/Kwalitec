#!/usr/bin/env python3
"""Generate P-003 Welcome Pack PDF (ReportLab) and DOCX from brand-locked content."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from reportlab.lib.colors import Color, HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader

ROOT = Path(__file__).resolve().parent
PDF_PATH = ROOT / "Kwalitec_Private_Beta_Welcome.pdf"
DOCX_PATH = ROOT / "Kwalitec_Private_Beta_Welcome.docx"
HTML_PATH = ROOT / "Kwalitec_Private_Beta_Welcome.html"
LOGO_DARK = ROOT / "assets" / "logo-primary-dark.png"
FONTS = ROOT / "assets" / "fonts"

# Brand tokens (COLOUR_SPECIFICATION.md)
NAVY = HexColor("#0A1628")
BLUE = HexColor("#3B4FB8")
GOLD = HexColor("#E8B02B")
IVORY = HexColor("#FAF8F2")
SECTION = HexColor("#F1F3F7")
TEXT = HexColor("#1E2430")
SECONDARY = HexColor("#4A5568")
MUTED_DARK = HexColor("#8B93A7")
CALLOUT_BG = HexColor("#FFFCF5")
WHITE = HexColor("#FFFFFF")
BORDER = HexColor("#E6E9EF")

PAGE_W, PAGE_H = A4  # 595.27 x 841.89 pt


def register_fonts() -> None:
    pdfmetrics.registerFont(TTFont("Inter", str(FONTS / "Inter-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("Inter-Medium", str(FONTS / "Inter-Medium.ttf")))
    pdfmetrics.registerFont(TTFont("Inter-SemiBold", str(FONTS / "Inter-SemiBold.ttf")))
    pdfmetrics.registerFont(TTFont("Inter-Bold", str(FONTS / "Inter-Bold.ttf")))


def wrap_text(c: canvas.Canvas, text: str, font: str, size: float, max_width: float) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = word if not current else f"{current} {word}"
        if c.stringWidth(trial, font, size) <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_wrapped(
    c: canvas.Canvas,
    text: str,
    x: float,
    y: float,
    font: str,
    size: float,
    max_width: float,
    leading: float,
    color: Color,
    align: str = "left",
) -> float:
    """Draw wrapped text; return y of last baseline."""
    c.setFillColor(color)
    c.setFont(font, size)
    lines = wrap_text(c, text, font, size, max_width)
    for i, line in enumerate(lines):
        yy = y - i * leading
        if align == "center":
            c.drawCentredString(x + max_width / 2, yy, line)
        elif align == "right":
            c.drawRightString(x + max_width, yy, line)
        else:
            c.drawString(x, yy, line)
    return y - (len(lines) - 1) * leading


def draw_rounded_rect(
    c: canvas.Canvas,
    x: float,
    y: float,
    w: float,
    h: float,
    r: float,
    fill: Color | None = None,
    stroke: Color | None = None,
    stroke_width: float = 0.75,
) -> None:
    c.saveState()
    if fill is not None:
        c.setFillColor(fill)
    if stroke is not None:
        c.setStrokeColor(stroke)
        c.setLineWidth(stroke_width)
    op = 1 if fill is not None else 0
    stroke_op = 1 if stroke is not None else 0
    c.roundRect(x, y, w, h, r, stroke=stroke_op, fill=op)
    c.restoreState()


def draw_lucide_book(c: canvas.Canvas, x: float, y: float, s: float, color: Color) -> None:
    c.saveState()
    c.setStrokeColor(color)
    c.setLineWidth(1.1)
    c.setLineCap(1)
    c.setLineJoin(1)
    # Simplified open-book outline
    p = c.beginPath()
    p.moveTo(x, y + s * 0.15)
    p.lineTo(x, y + s * 0.85)
    p.curveTo(x, y + s, x + s * 0.15, y + s, x + s * 0.35, y + s * 0.9)
    p.lineTo(x + s * 0.5, y + s * 0.78)
    p.lineTo(x + s * 0.65, y + s * 0.9)
    p.curveTo(x + s * 0.85, y + s, x + s, y + s, x + s, y + s * 0.85)
    p.lineTo(x + s, y + s * 0.15)
    c.drawPath(p, stroke=1, fill=0)
    c.line(x + s * 0.5, y + s * 0.2, x + s * 0.5, y + s * 0.78)
    c.restoreState()


def draw_lucide_check(c: canvas.Canvas, x: float, y: float, s: float, color: Color) -> None:
    c.saveState()
    c.setStrokeColor(color)
    c.setLineWidth(1.1)
    c.setLineCap(1)
    c.circle(x + s / 2, y + s / 2, s * 0.42, stroke=1, fill=0)
    p = c.beginPath()
    p.moveTo(x + s * 0.28, y + s * 0.5)
    p.lineTo(x + s * 0.44, y + s * 0.34)
    p.lineTo(x + s * 0.72, y + s * 0.66)
    c.drawPath(p, stroke=1, fill=0)
    c.restoreState()


def draw_lucide_message(c: canvas.Canvas, x: float, y: float, s: float, color: Color) -> None:
    c.saveState()
    c.setStrokeColor(color)
    c.setLineWidth(1.1)
    c.setLineCap(1)
    c.setLineJoin(1)
    c.roundRect(x + s * 0.12, y + s * 0.32, s * 0.76, s * 0.5, s * 0.08, stroke=1, fill=0)
    p = c.beginPath()
    p.moveTo(x + s * 0.28, y + s * 0.32)
    p.lineTo(x + s * 0.22, y + s * 0.12)
    p.lineTo(x + s * 0.48, y + s * 0.32)
    c.drawPath(p, stroke=1, fill=0)
    c.restoreState()


def draw_lucide_alert(c: canvas.Canvas, x: float, y: float, s: float, color: Color) -> None:
    c.saveState()
    c.setStrokeColor(color)
    c.setLineWidth(1.1)
    c.setLineCap(1)
    c.circle(x + s / 2, y + s / 2, s * 0.42, stroke=1, fill=0)
    c.line(x + s / 2, y + s * 0.32, x + s / 2, y + s * 0.55)
    c.circle(x + s / 2, y + s * 0.68, 0.7, stroke=0, fill=1)
    c.setFillColor(color)
    c.circle(x + s / 2, y + s * 0.68, 0.85, stroke=0, fill=1)
    c.restoreState()


def build_pdf() -> None:
    register_fonts()
    c = canvas.Canvas(str(PDF_PATH), pagesize=A4)
    c.setTitle("Kwalitec Private Beta — Welcome Pack")
    c.setAuthor("Kwalitec")
    c.setSubject("Founding Cohort Welcome Pack 2026")
    c.setKeywords("Kwalitec, Private Beta, Founding Cohort")
    c.setCreator("Kwalitec P-003 Welcome Pack")

    margin = 0  # full-bleed ivory within gold border
    border = 1  # ~1pt

    # Page ivory fill
    c.setFillColor(IVORY)
    c.rect(0, 0, PAGE_W, PAGE_H, stroke=0, fill=1)

    # Gold page border
    c.setStrokeColor(GOLD)
    c.setLineWidth(1)
    inset = 4
    c.rect(inset, inset, PAGE_W - 2 * inset, PAGE_H - 2 * inset, stroke=1, fill=0)

    # ——— Header (~top 23–25%) ———
    header_h = 68 * mm
    header_bottom = PAGE_H - inset - header_h
    c.setFillColor(NAVY)
    c.rect(inset, header_bottom, PAGE_W - 2 * inset, header_h, stroke=0, fill=1)

    # Logo top-right, ~42mm wide — transparent white lockup for navy
    content_pad = 12 * mm
    logo_w = 42 * mm
    # Official lockup aspect from approved transparent master (~1418×371)
    logo_h = logo_w * (371 / 1418)
    logo_x = PAGE_W - content_pad - logo_w
    logo_y = PAGE_H - inset - 8 * mm - logo_h
    c.drawImage(
        ImageReader(str(LOGO_DARK)),
        logo_x,
        logo_y,
        width=logo_w,
        height=logo_h,
        mask="auto",
        preserveAspectRatio=True,
        anchor="c",
    )

    # Title
    title_x = content_pad
    title_y = header_bottom + 26 * mm
    c.setFillColor(WHITE)
    c.setFont("Inter-Bold", 24)
    c.drawString(title_x, title_y + 12, "Welcome to")
    c.drawString(title_x, title_y - 10, "Kwalitec Private Beta")

    # Meta
    c.setFillColor(MUTED_DARK)
    c.setFont("Inter-Medium", 8.5)
    meta = "FOUNDING COHORT"
    c.drawString(title_x, header_bottom + 9 * mm, meta)
    mx = title_x + c.stringWidth(meta, "Inter-Medium", 8.5) + 6
    c.setFillColor(GOLD)
    c.drawString(mx, header_bottom + 9 * mm, "·")
    mx += 8
    c.setFillColor(MUTED_DARK)
    c.drawString(mx, header_bottom + 9 * mm, "PRIVATE BETA")
    mx += c.stringWidth("PRIVATE BETA", "Inter-Medium", 8.5) + 6
    c.setFillColor(GOLD)
    c.drawString(mx, header_bottom + 9 * mm, "·")
    mx += 8
    c.setFillColor(MUTED_DARK)
    c.drawString(mx, header_bottom + 9 * mm, "2026")

    # Gold divider under header
    c.setStrokeColor(GOLD)
    c.setLineWidth(1.25)
    c.line(inset, header_bottom, PAGE_W - inset, header_bottom)

    # ——— Body ———
    y = header_bottom - 6.5 * mm
    content_w = PAGE_W - 2 * content_pad

    def section_heading(label: str, yy: float) -> float:
        c.setFillColor(BLUE)
        c.setFont("Inter-SemiBold", 10.5)
        c.drawString(content_pad, yy, label)
        return yy - 4.0 * mm

    def body_para(text: str, yy: float, size: float = 9.5, leading: float = 12.6) -> float:
        last = draw_wrapped(c, text, content_pad, yy, "Inter", size, content_w, leading, TEXT)
        return last - 4.8 * mm

    # Welcome
    y = section_heading("Welcome", y)
    y = body_para(
        "Thank you for joining the Kwalitec Private Beta. You are among the very first students "
        "invited into this Founding Cohort. Your participation will help shape how Kwalitec supports "
        "actuarial learners preparing for professional examinations. We value your time, curiosity, "
        "and honest feedback as we refine the product with care.",
        y,
    )

    # What is Kwalitec?
    y = section_heading("What is Kwalitec?", y)
    y = body_para(
        "Kwalitec is an AI-powered study companion. It complements official IFoA study resources and "
        "is designed to bring clarity to what to study next and why it matters. It does not replace "
        "your CMP, Core Reading, tutorials, or professional judgement. Those remain the foundation of "
        "your preparation. Kwalitec is here to guide study decisions with honesty — not to substitute "
        "the work of becoming a professional.",
        y,
    )

    # Your Role
    y = section_heading("Your Role", y)
    roles = [
        (draw_lucide_book, "Study normally using Kwalitec alongside your official resources."),
        (draw_lucide_check, "Complete your Daily Missions and Study Sessions."),
        (draw_lucide_message, "Share honest feedback and suggestions."),
        (draw_lucide_alert, "Report issues to help improve the product before public release."),
    ]
    gap = 2.4 * mm
    col_w = (content_w - gap) / 2
    row_h = 13.5 * mm
    icon_s = 5 * mm
    for i, (icon_fn, text) in enumerate(roles):
        col = i % 2
        row = i // 2
        rx = content_pad + col * (col_w + gap)
        ry = y - (row + 1) * row_h - row * gap + 2 * mm
        draw_rounded_rect(c, rx, ry, col_w, row_h, 2.2, fill=SECTION, stroke=BORDER, stroke_width=0.5)
        icon_fn(c, rx + 2.4 * mm, ry + row_h / 2 - icon_s / 2, icon_s, BLUE)
        text_x = rx + 2.4 * mm + icon_s + 2.2 * mm
        text_w = col_w - (text_x - rx) - 2.4 * mm
        lines = wrap_text(c, text, "Inter", 8.4, text_w)
        text_block_h = len(lines) * 10.6
        ty = ry + (row_h + text_block_h) / 2 - 8.4
        draw_wrapped(c, text, text_x, ty, "Inter", 8.4, text_w, 10.6, TEXT)

    y = y - 2 * row_h - gap - 5 * mm

    # Please note callout
    callout_h = 28 * mm
    callout_y = y - callout_h
    draw_rounded_rect(
        c, content_pad, callout_y, content_w, callout_h, 2.2, fill=CALLOUT_BG, stroke=GOLD, stroke_width=1
    )
    c.setFillColor(NAVY)
    c.setFont("Inter-SemiBold", 8)
    c.drawString(content_pad + 4 * mm, callout_y + callout_h - 5.5 * mm, "PLEASE NOTE")
    note1 = (
        "This is an invite-only Private Beta. Kwalitec is designed to support your learning journey. "
        "It does not replace official study material. It does not guarantee examination success."
    )
    note2 = "Success will always depend on your own preparation, practice and commitment."
    ny = callout_y + callout_h - 10 * mm
    ny = draw_wrapped(
        c, note1, content_pad + 4 * mm, ny, "Inter", 8.6, content_w - 8 * mm, 11.2, SECONDARY
    )
    draw_wrapped(
        c, note2, content_pad + 4 * mm, ny - 3.2 * mm, "Inter", 8.6, content_w - 8 * mm, 11.2, SECONDARY
    )
    y = callout_y - 5.5 * mm

    # Getting Started + Our Commitment
    left_w = content_w * 0.56
    right_w = content_w * 0.40
    right_x = content_pad + content_w - right_w

    c.setFillColor(BLUE)
    c.setFont("Inter-SemiBold", 10.5)
    c.drawString(content_pad, y, "Getting Started")

    steps = [
        ("Visit ", "https://kwalitec.onrender.com"),
        ("Sign in using the credentials provided in your invitation email.", None),
        ("Change your temporary password.", None),
        ("Complete your first Daily Mission.", None),
        ("Share feedback through the application.", None),
    ]

    step_y = y - 5 * mm
    for i, (label, link) in enumerate(steps, start=1):
        cx = content_pad + 2.2 * mm
        cy = step_y + 1.2
        c.setFillColor(BLUE)
        c.circle(cx, cy, 2.1 * mm, stroke=0, fill=1)
        c.setFillColor(WHITE)
        c.setFont("Inter-SemiBold", 7)
        c.drawCentredString(cx, cy - 2.2, str(i))

        tx = content_pad + 6 * mm
        tw = left_w - 8 * mm
        if link:
            c.setFillColor(TEXT)
            c.setFont("Inter", 8.6)
            c.drawString(tx, step_y - 0.5, label)
            link_x = tx + c.stringWidth(label, "Inter", 8.6)
            c.setFillColor(BLUE)
            c.setFont("Inter-Medium", 8.6)
            c.drawString(link_x, step_y - 0.5, link)
            c.linkURL(
                link,
                (
                    link_x - 1,
                    step_y - 2,
                    link_x + c.stringWidth(link, "Inter-Medium", 8.6) + 1,
                    step_y + 9,
                ),
                relative=0,
            )
            c.setStrokeColor(BLUE)
            c.setLineWidth(0.6)
            c.line(
                link_x,
                step_y - 1.5,
                link_x + c.stringWidth(link, "Inter-Medium", 8.6),
                step_y - 1.5,
            )
            step_y -= 5.4 * mm
        else:
            last = draw_wrapped(c, label, tx, step_y, "Inter", 8.6, tw, 11, TEXT)
            step_y = last - 4.6 * mm

    # Commitment panel (heading inside)
    commit_h = 38 * mm
    commit_bottom = y - commit_h + 3 * mm
    draw_rounded_rect(
        c, right_x, commit_bottom, right_w, commit_h, 2.2, fill=SECTION, stroke=BORDER, stroke_width=0.5
    )
    cy = commit_bottom + commit_h - 5.5 * mm
    c.setFillColor(BLUE)
    c.setFont("Inter-SemiBold", 10.5)
    c.drawString(right_x + 3 * mm, cy, "Our Commitment")
    cy -= 4.5 * mm
    cy = draw_wrapped(
        c,
        "Every improvement to Kwalitec is measured against one question:",
        right_x + 3 * mm,
        cy,
        "Inter",
        8.2,
        right_w - 6 * mm,
        10.8,
        SECONDARY,
    )
    cy -= 3.0 * mm
    cy = draw_wrapped(
        c,
        "“Does this help students become better actuarial professionals?”",
        right_x + 3 * mm,
        cy,
        "Inter-Medium",
        8.4,
        right_w - 6 * mm,
        11,
        TEXT,
    )
    cy -= 3.0 * mm
    draw_wrapped(
        c,
        "As a Founding Cohort member, your feedback will directly influence the future of Kwalitec.",
        right_x + 3 * mm,
        cy,
        "Inter",
        8.2,
        right_w - 6 * mm,
        10.8,
        SECONDARY,
    )

    # Footer — sit just below the lower content block
    content_bottom = min(step_y, commit_bottom) - 6 * mm
    footer_y = max(inset + 8 * mm, content_bottom - 2 * mm)
    c.setStrokeColor(GOLD)
    c.setLineWidth(1)
    c.line(content_pad, footer_y + 5 * mm, PAGE_W - content_pad, footer_y + 5 * mm)

    c.setFillColor(SECONDARY)
    c.setFont("Inter-Medium", 7)
    left_meta = "PRIVATE BETA"
    c.drawString(content_pad, footer_y, left_meta)
    mx = content_pad + c.stringWidth(left_meta, "Inter-Medium", 7) + 5
    c.setFillColor(GOLD)
    c.drawString(mx, footer_y, "·")
    mx += 7
    c.setFillColor(SECONDARY)
    mid = "Version 2.0.0-beta.1"
    c.drawString(mx, footer_y, mid)
    mx += c.stringWidth(mid, "Inter-Medium", 7) + 5
    c.setFillColor(GOLD)
    c.drawString(mx, footer_y, "·")
    mx += 7
    c.setFillColor(SECONDARY)
    c.drawString(mx, footer_y, "INVITE ONLY")

    c.setFillColor(NAVY)
    c.setFont("Inter-Medium", 7.5)
    c.drawRightString(PAGE_W - content_pad, footer_y, "Helping students become professionals.")

    c.showPage()
    c.save()


def set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run_font(run, size_pt: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Inter"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Inter")
    rFonts.set(qn("w:hAnsi"), "Inter")
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "3B4FB8")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "19")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Inter")
    rFonts.set(qn("w:hAnsi"), "Inter")
    rPr.extend([rFonts, color_el, u, sz])
    new_run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.set("{http://www.w3.org and}xml:space" if False else qn("xml:space"), "preserve")
    # fix namespace for xml:space
    text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_page_border(section) -> None:
    sectPr = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "12")
        border.set(qn("w:color"), "E8B02B")
        pgBorders.append(border)
    sectPr.append(pgBorders)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    section.left_margin = Mm(12)
    section.right_margin = Mm(12)
    set_page_border(section)

    header_table = doc.add_table(rows=1, cols=1)
    cell = header_table.cell(0, 0)
    set_cell_shading(cell, "0A1628")

    p_logo = cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_logo.add_run()
    run.add_picture(str(LOGO_DARK), width=Mm(42))

    p_title = cell.add_paragraph()
    r = p_title.add_run("Welcome to\nKwalitec Private Beta")
    set_run_font(r, 20, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    p_meta = cell.add_paragraph()
    r = p_meta.add_run("FOUNDING COHORT  ·  PRIVATE BETA  ·  2026")
    set_run_font(r, 8, color=RGBColor(0x8B, 0x93, 0xA7))

    def heading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        set_run_font(r, 11, bold=True, color=RGBColor(0x3B, 0x4F, 0xB8))

    def body(text: str, size: float = 10) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.35
        r = p.add_run(text)
        set_run_font(r, size, color=RGBColor(0x1E, 0x24, 0x30))

    heading("Welcome")
    body(
        "Thank you for joining the Kwalitec Private Beta. You are among the very first students "
        "invited into this Founding Cohort. Your participation will help shape how Kwalitec supports "
        "actuarial learners preparing for professional examinations. We value your time, curiosity, "
        "and honest feedback as we refine the product with care."
    )

    heading("What is Kwalitec?")
    body(
        "Kwalitec is an AI-powered study companion. It complements official IFoA study resources and "
        "is designed to bring clarity to what to study next and why it matters. It does not replace "
        "your CMP, Core Reading, tutorials, or professional judgement. Those remain the foundation of "
        "your preparation. Kwalitec is here to guide study decisions with honesty — not to substitute "
        "the work of becoming a professional."
    )

    heading("Your Role")
    for item in [
        "Study normally using Kwalitec alongside your official resources.",
        "Complete your Daily Missions and Study Sessions.",
        "Share honest feedback and suggestions.",
        "Report issues to help improve the product before public release.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_run_font(r, 9.5, color=RGBColor(0x1E, 0x24, 0x30))

    heading("Please note")
    body(
        "This is an invite-only Private Beta. Kwalitec is designed to support your learning journey. "
        "It does not replace official study material. It does not guarantee examination success. "
        "Success will always depend on your own preparation, practice and commitment.",
        size=9.5,
    )

    heading("Getting Started")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("1. Visit ")
    set_run_font(r, 9.5, color=RGBColor(0x1E, 0x24, 0x30))
    add_hyperlink(p, "https://kwalitec.onrender.com", "https://kwalitec.onrender.com")

    for i, text in enumerate(
        [
            "Sign in using the credentials provided in your invitation email.",
            "Change your temporary password.",
            "Complete your first Daily Mission.",
            "Share feedback through the application.",
        ],
        start=2,
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{i}. {text}")
        set_run_font(r, 9.5, color=RGBColor(0x1E, 0x24, 0x30))

    heading("Our Commitment")
    body(
        "Every improvement to Kwalitec is measured against one question: "
        "“Does this help students become better actuarial professionals?” "
        "As a Founding Cohort member, your feedback will directly influence the future of Kwalitec.",
        size=9.5,
    )

    p_foot = doc.add_paragraph()
    pPr = p_foot._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "10")
    top.set(qn("w:space"), "10")
    top.set(qn("w:color"), "E8B02B")
    pBdr.append(top)
    pPr.append(pBdr)
    r = p_foot.add_run(
        "PRIVATE BETA  ·  Version 2.0.0-beta.1  ·  Invite Only\n"
        "Helping students become professionals."
    )
    set_run_font(r, 8, color=RGBColor(0x4A, 0x55, 0x68))

    core = doc.core_properties
    core.title = "Kwalitec Private Beta — Welcome Pack"
    core.author = "Kwalitec"
    core.subject = "Founding Cohort Welcome Pack 2026"
    core.keywords = "Kwalitec; Private Beta; Founding Cohort"

    doc.save(str(DOCX_PATH))


def main() -> None:
    assert HTML_PATH.exists(), "HTML master missing"
    build_pdf()
    build_docx()
    # Verify single page
    from pypdf import PdfReader  # optional

    try:
        n = len(PdfReader(str(PDF_PATH)).pages)
    except Exception:
        # fallback: reportlab always wrote one showPage
        n = 1
    print(f"Wrote {PDF_PATH} (pages={n})")
    print(f"Wrote {DOCX_PATH}")
    print(f"HTML master: {HTML_PATH}")


if __name__ == "__main__":
    main()
